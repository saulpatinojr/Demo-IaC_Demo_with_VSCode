from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = str(
    Path(__file__).resolve().parent / "IaC_Demo_Hourly_Cost_Printout.docx"
)

BLUE = "2E74B5"
DARK = "1F4D78"
LIGHT = "E8EEF5"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shade_elem = OxmlElement("w:shd")
    shade_elem.set(qn("w:fill"), fill)
    tc_pr.append(shade_elem)


def margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, value in (
        ("top", 100),
        ("start", 120),
        ("bottom", 100),
        ("end", 120),
    ):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tc_pr.append(mar)


def font(run, size=11, color="222222", bold=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def para(doc, text="", after=6, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if text:
        font(p.add_run(text))
    return p


def head(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    font(
        p.add_run(text),
        16 if level == 1 else 13,
        BLUE if level < 3 else DARK,
        True,
    )
    return p


def table(doc, headers, rows, widths):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False

    for i, header in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        margins(cell)
        shade(cell, LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(header), 9.5, DARK, True)

    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            font(p.add_run(value), 9.5)
    return tbl


doc = Document()
sec = doc.sections[0]
for margin_name in (
    "top_margin",
    "bottom_margin",
    "left_margin",
    "right_margin",
):
    setattr(sec, margin_name, Inches(1))

sec.header_distance = sec.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
):
    style = doc.styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(3)
font(title.add_run("IaC Demo Resource Cost Printout"), 24, "0B2545", True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
font(
    subtitle.add_run("Estimated average cost for one complete running hour"),
    15,
    BLUE,
)

para(
    doc,
    (
        "Scope: Azure resources defined in the Demo-IaC repository. "
        "Levels are cumulative: L2 includes L1, L3 includes the earlier "
        "foundation, and L4 adds the second-region and global-access "
        "components."
    ),
)

head(doc, "Executive summary")
table(
    doc,
    [
        "Level",
        "Incremental cost added by level",
        "Estimated total cost for one hour",
    ],
    [
        ("L1 - Hub and spoke", "$0.25/hr", "$0.25/hr"),
        ("L2 - Web tier and firewall", "$0.57/hr", "$0.82/hr"),
        ("L3 - Containers and data", "$0.08/hr", "$0.90/hr"),
        ("L4 - Global scale", "$0.06/hr", "$0.96/hr"),
    ],
    [1.65, 2.25, 2.60],
)

para(
    doc,
    (
        "Planning figure: approximately $1.00 for one hour with the full "
        "L4 stack running, before traffic, log ingestion, backups, or "
        "unusual data transfer. Actual billing may vary by subscription "
        "agreement, region, currency, VM price, and usage."
    ),
)

head(doc, "Cost progression by level")
for title_text, body in (
    (
        "L1 - Hub and spoke | approximately $0.25 per hour",
        (
            "L1 contains one Standard_B2s Linux VM, one Basic Bastion host, "
            "a Standard public IP used by Bastion, and two virtual networks "
            "with peering. VNets and NSGs are generally not hourly charges. "
            "The VM and Bastion are the main always-on costs."
        ),
    ),
    (
        "L2 - Web tier and firewall | approximately $0.82 per hour total",
        (
            "L2 retains L1 and adds three Standard_B2s Linux VMs, three "
            "operating-system disks, an internal Standard Load Balancer, an "
            "Azure Firewall Standard instance, and the firewall public IP. "
            "The firewall is the dominant new cost; the three VMs are the "
            "next largest addition."
        ),
    ),
    (
        "L3 - Containers and data | approximately $0.90 per hour total",
        (
            "L3 adds a Container Apps environment and one minimum-size "
            "container replica, Azure SQL Database Basic, two private "
            "endpoints, Key Vault, private DNS, Log Analytics, Application "
            "Insights, a managed identity, and alerting. The estimate assumes "
            "very light demo traffic and negligible log ingestion."
        ),
    ),
    (
        "L4 - Global scale | approximately $0.96 per hour total",
        (
            "L4 adds a second-region Container Apps environment and one "
            "minimum-size replica, a secondary SQL server/failover "
            "configuration, and Azure Front Door Standard. Front Door has a "
            "base hourly charge; requests and data transfer are additional "
            "usage meters."
        ),
    ),
):
    head(doc, title_text, 2)
    para(doc, body)

head(doc, "Illustrative resource breakdown")
table(
    doc,
    ["Resource group / level", "Resources counted", "Approx. hourly amount"],
    [
        ("L1", "1 B2s VM + disk; Basic Bastion; Bastion public IP", "$0.25"),
        (
            "L2 incremental",
            "3 B2s VMs + disks; Azure Firewall Standard; firewall public IP; "
            "Standard Load Balancer",
            "$0.57",
        ),
        (
            "L3 incremental",
            "Container Apps; Basic SQL database; private endpoints; "
            "monitoring and platform services",
            "$0.08",
        ),
        (
            "L4 incremental",
            "Secondary Container App; Front Door Standard; failover-related "
            "services",
            "$0.06",
        ),
        ("Full L4 stack", "All resources from L1 through L4", "$0.96"),
    ],
    [1.50, 3.65, 1.35],
)

head(doc, "Assumptions behind the estimate")
for bullet in (
    (
        "Region: East US 2 for L1-L3, with West US 2 as the L4 secondary "
        "region, matching the Bicep defaults."
    ),
    (
        "Pricing basis: approximate US pay-as-you-go list pricing, "
        "rounded for presentation use."
    ),
    (
        "VMs run for the entire hour and are not deallocated. Standard_LRS "
        "OS disks are included as a small hourly equivalent."
    ),
    (
        "Container Apps run at the configured minimum of one 0.5 vCPU / "
        "1 GiB replica per region."
    ),
    (
        "L3 assumes a Basic 5-DTU Azure SQL database and minimal data "
        "storage, backups, requests, and log ingestion."
    ),
    (
        "Network traffic, NAT/data processing, SQL backup growth, Log "
        "Analytics ingestion, Application Insights telemetry, and Front Door "
        "requests/data transfer are excluded or assumed negligible."
    ),
    (
        "The optional L1 VPN Gateway is not included because it is "
        "commented out in the template."
    ),
):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    font(p.add_run(bullet))

head(doc, "Important cost notes")
para(
    doc,
    (
        "Azure Firewall and Bastion continue billing while deployed, even if "
        "the demo is idle. The repository includes a teardown workflow; use "
        "it after the demonstration or when the environment is not needed. "
        "For a live presentation, budget a small buffer above the table, "
        "such as $1.25-$1.50 for a full L4 hour, to cover normal request and "
        "telemetry activity."
    ),
)
para(
    doc,
    (
        "These are planning estimates, not an invoice or quote. Azure pricing "
        "pages state that displayed prices can vary by agreement, date, "
        "currency, and region. Confirm the final amount with the Azure "
        "Pricing Calculator or the subscription Cost Management view before "
        "presenting a committed budget."
    ),
)

head(doc, "Reference links")
for label, url in (
    ("Azure pricing overview", "https://azure.microsoft.com/pricing/"),
    (
        "Azure Firewall pricing",
        "https://azure.microsoft.com/pricing/details/azure-firewall/",
    ),
    (
        "Azure Bastion pricing",
        "https://azure.microsoft.com/pricing/details/azure-bastion/",
    ),
    (
        "Azure SQL Database pricing",
        (
            "https://azure.microsoft.com/pricing/details/"
            "azure-sql-database/single/"
        ),
    ),
    (
        "Azure Front Door pricing",
        "https://azure.microsoft.com/pricing/details/frontdoor/",
    ),
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{label}: {url}")
    font(run, 9, "555555")

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(footer.add_run("IaC Demo | Cost Printout"), 9, "777777")

doc.core_properties.title = "IaC Demo Resource Cost Printout"
doc.core_properties.subject = (
    "Approximate one-hour Azure resource cost by cumulative lab level"
)

doc.save(OUT)
print(OUT)

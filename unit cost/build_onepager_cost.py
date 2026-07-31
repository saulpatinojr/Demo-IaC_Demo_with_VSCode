from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
OUT = str(Path(__file__).resolve().parent / 'IaC_Demo_Cost_One_Page.docx'); BLUE='2E74B5'; DARK='1F4D78'; LIGHT='E8EEF5'
def font(r,size=9,color='222222',bold=False):
    r.font.name='Calibri'; r._element.get_or_add_rPr().rFonts.set(qn('w:ascii'),'Calibri'); r._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'),'Calibri'); r.font.size=Pt(size); r.bold=bold; r.font.color.rgb=RGBColor.from_string(color)
def shade(c,fill):
    x=c._tc.get_or_add_tcPr(); s=OxmlElement('w:shd'); s.set(qn('w:fill'),fill); x.append(s)
def cellm(c):
    x=c._tc.get_or_add_tcPr(); m=OxmlElement('w:tcMar')
    for side,val in [('top',60),('start',100),('bottom',60),('end',100)]:
        n=OxmlElement('w:'+side); n.set(qn('w:w'),str(val)); n.set(qn('w:type'),'dxa'); m.append(n)
    x.append(m)
def para(d,s='',after=2,before=0,size=9,color='222222',bold=False):
    p=d.add_paragraph(); p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.0
    if s: font(p.add_run(s),size,color,bold)
    return p
def level(d,title,summary,resources,price,total,nextstep):
    p=para(d,title,1,4,11.2,BLUE,True); para(d,summary,2,0,8.5)
    t=d.add_table(rows=1,cols=3); t.autofit=False; widths=[2.7,1.15,2.65]
    for i,h in enumerate(['Main resources','$/hour','Summed cost']):
        c=t.rows[0].cells[i]; c.width=Inches(widths[i]); shade(c,LIGHT); cellm(c); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(h),8.2,DARK,True)
    cs=t.add_row().cells
    for i,v in enumerate([resources,price,total]):
        cs[i].width=Inches(widths[i]); cellm(cs[i]); p=cs[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0; font(p.add_run(v if i<2 else total),8.1)
    para(d,'Builds next: '+nextstep,2,0,8.2,DARK,True)
d=Document(); sec=d.sections[0]
for a in ['top_margin','bottom_margin','left_margin','right_margin']: setattr(sec,a,Inches(.55))
sec.header_distance=sec.footer_distance=Inches(.25); s=d.styles['Normal']; s.font.name='Calibri'; s.font.size=Pt(9); s.paragraph_format.space_after=Pt(2); s.paragraph_format.line_spacing=1.0
p=para(d,'IaC Demo Azure Cost Progression',1,0,19,'0B2545',True); para(d,'One-page resource cost handout | average cost for one complete running hour',4,0,9.5,BLUE,True)
para(d,'Assumptions: East US 2 for L1-L3 and West US 2 for the L4 secondary region. Approximate US pay-as-you-go list-price estimates, rounded for presentation. Levels are cumulative; traffic, data transfer, logging, backups, and unusual usage are excluded or assumed minimal.',5,0,8.1,'555555')
level(d,'L1 - Hub and spoke | total: ~$0.25/hour','Creates the secure connectivity foundation and a test workload.','1 B2s VM + disk','VM/disk ~$0.05; Bastion ~$0.19; public IP ~$0.01; VNets/peering ~$0.00 fixed','$0.25/hour','L2 reuses the hub/spoke networks and adds a protected, load-balanced web tier and firewall.')
level(d,'L2 - Web tier and firewall | adds ~$0.57/hour','Adds redundancy and traffic governance to the L1 network.','3 B2s VMs + disks; Firewall; public IP; Load Balancer','VMs/disks ~$0.15; Firewall ~$0.39; public IP ~$0.01; LB ~$0.02; NSGs/routes ~$0.00 fixed','$0.82/hour cumulative','L3 adds containers, private data services, identity, and monitoring on the existing foundation.')
level(d,'L3 - Containers and data | adds ~$0.08/hour','Moves the demo toward a modern, secure application platform.','Container Apps; SQL Basic; private endpoints; Key Vault; monitoring','ACA ~$0.02; SQL ~$0.01; private endpoints ~$0.02; other platform services ~$0.03','$0.90/hour cumulative','L4 adds a second-region app, SQL failover capability, and global Front Door.')
level(d,'L4 - Global scale | adds ~$0.06/hour','Adds regional resilience and a single global entry point.','Secondary Container App; secondary SQL/failover; Front Door','Secondary ACA ~$0.01; SQL/failover ~$0.01; Front Door base ~$0.05; traffic variable','$0.96/hour cumulative','End state: multi-region app delivery with health-probed routing and database failover.')
para(d,'Presenter takeaway: budget approximately $1.25-$1.50 for one full hour of the L4 demo to allow for normal telemetry and traffic. Firewall and Bastion bill while deployed, even when idle; run teardown when finished.',3,3,8.1,'7A5A00',True)
para(d,'Planning estimates only. Confirm final prices in the Azure Pricing Calculator or Cost Management for the target subscription. Reference: azure.microsoft.com/pricing',0,0,7.4,'666666')
f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(f.add_run('IaC Demo | One-page cost handout'),7.2,'777777')
d.core_properties.title='IaC Demo Azure Cost Progression - One Page'; d.save(OUT); print(OUT)
